"""
Extract structured content from PDF, DOCX, and XLSX files.

Reads source documents and returns normalized content elements (headings,
paragraphs, tables, lists) that can be converted into EditorJS blocks for
the Rattle API.

Usage:
    python tools/extract_content.py document.pdf
    python tools/extract_content.py document.docx --format editorjs
    python tools/extract_content.py spreadsheet.xlsx --sheet "Specs"
    python tools/extract_content.py document.pdf -o output.json

Supported formats:
    .pdf   — extracts text blocks and tables via PyMuPDF
    .docx  — extracts paragraphs, headings, tables, lists via python-docx
    .xlsx  — extracts rows as table data via openpyxl

Output formats:
    json      — normalized content elements (default)
    editorjs  — EditorJS blocks ready for the Rattle API

Requirements:
    pip install PyMuPDF python-docx openpyxl
"""

import argparse
import json
import os
import re
import sys
import uuid

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ---------------------------------------------------------------------------
# Normalized content elements
# ---------------------------------------------------------------------------
# All extractors produce a list of these dicts:
#
#   {"type": "heading",   "level": int, "text": str}
#   {"type": "paragraph", "text": str}
#   {"type": "table",     "headers": [str], "rows": [[str]]}
#   {"type": "list",      "style": "unordered"|"ordered", "items": [str]}


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------
def extract_pdf(filepath):
    """Extract structured content from a PDF file.

    Uses PyMuPDF to read text blocks and tables. Detects headings based on
    font size relative to the document's body text.
    """
    import fitz

    doc = fitz.open(filepath)
    elements = []

    # First pass: determine median font size (= body text size)
    all_sizes = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:  # text blocks only
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"].strip()
                    if text:
                        all_sizes.append(span["size"])

    if not all_sizes:
        doc.close()
        return elements

    all_sizes.sort()
    body_size = all_sizes[len(all_sizes) // 2]
    heading_threshold = body_size * 1.15

    # Second pass: extract content
    for page_idx in range(len(doc)):
        page = doc[page_idx]

        # Try extracting tables first (PyMuPDF >= 1.23.0)
        try:
            tables = page.find_tables()
            table_rects = []
            for table in tables:
                table_rects.append(table.bbox)
                data = table.extract()
                if not data or len(data) < 2:
                    continue
                headers = [str(c) if c else "" for c in data[0]]
                rows = []
                for row in data[1:]:
                    rows.append([str(c) if c else "" for c in row])
                if any(h.strip() for h in headers):
                    elements.append({
                        "type": "table",
                        "headers": headers,
                        "rows": rows,
                    })
        except AttributeError:
            # Older PyMuPDF without find_tables
            table_rects = []

        # Extract text blocks, skipping those inside table regions
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:
                continue

            # Skip blocks that overlap with detected tables
            bx0, by0, bx1, by1 = block["bbox"]
            in_table = False
            for tr in table_rects:
                tx0, ty0, tx1, ty1 = tr
                if bx0 >= tx0 - 5 and by0 >= ty0 - 5 and bx1 <= tx1 + 5 and by1 <= ty1 + 5:
                    in_table = True
                    break
            if in_table:
                continue

            # Collect spans from this block
            block_text_parts = []
            block_max_size = 0
            is_bold = False

            for line in block["lines"]:
                line_text = ""
                for span in line["spans"]:
                    text = span["text"].strip()
                    if text:
                        line_text += span["text"]
                        block_max_size = max(block_max_size, span["size"])
                        if "bold" in span["font"].lower() or span.get("flags", 0) & 2**4:
                            is_bold = True
                if line_text.strip():
                    block_text_parts.append(line_text.strip())

            full_text = " ".join(block_text_parts).strip()
            if not full_text:
                continue

            # Classify: heading or paragraph
            if block_max_size >= heading_threshold and len(full_text) < 200:
                level = 2 if block_max_size >= body_size * 1.5 else 3
                elements.append({"type": "heading", "level": level, "text": full_text})
            elif is_bold and len(full_text) < 150 and len(block_text_parts) == 1:
                elements.append({"type": "heading", "level": 3, "text": full_text})
            else:
                elements.append({"type": "paragraph", "text": full_text})

    doc.close()
    return _merge_paragraphs(elements)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def extract_docx(filepath):
    """Extract structured content from a DOCX file.

    Uses python-docx to read paragraphs (with style detection for headings
    and lists) and tables.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(filepath)
    elements = []
    current_list_items = []
    current_list_style = None

    def _flush_list():
        nonlocal current_list_items, current_list_style
        if current_list_items:
            elements.append({
                "type": "list",
                "style": current_list_style or "unordered",
                "items": current_list_items,
            })
            current_list_items = []
            current_list_style = None

    for element in doc.element.body:
        # Handle tables
        if element.tag.endswith('}tbl'):
            _flush_list()
            for table in doc.tables:
                if table._element is element:
                    rows_data = []
                    for row in table.rows:
                        rows_data.append([cell.text.strip() for cell in row.cells])
                    if len(rows_data) >= 2:
                        elements.append({
                            "type": "table",
                            "headers": rows_data[0],
                            "rows": rows_data[1:],
                        })
                    elif rows_data:
                        elements.append({
                            "type": "table",
                            "headers": rows_data[0],
                            "rows": [],
                        })
                    break

        # Handle paragraphs
        elif element.tag.endswith('}p'):
            for para in doc.paragraphs:
                if para._element is element:
                    text = para.text.strip()
                    if not text:
                        _flush_list()
                        break

                    style_name = (para.style.name or "").lower()

                    # Detect headings
                    if style_name.startswith("heading"):
                        _flush_list()
                        match = re.search(r'\d+', style_name)
                        level = int(match.group()) if match else 3
                        elements.append({"type": "heading", "level": level, "text": text})

                    # Detect lists
                    elif style_name.startswith("list") or _has_numbering(para):
                        if "bullet" in style_name or "list bullet" in style_name:
                            list_style = "unordered"
                        elif "number" in style_name or "list number" in style_name:
                            list_style = "ordered"
                        else:
                            list_style = "unordered"

                        if current_list_style and current_list_style != list_style:
                            _flush_list()
                        current_list_style = list_style
                        current_list_items.append(text)

                    # Regular paragraph
                    else:
                        _flush_list()
                        elements.append({"type": "paragraph", "text": text})

                    break

    _flush_list()
    return elements


def _has_numbering(para):
    """Check if a DOCX paragraph has numbering (bullet or numbered list)."""
    pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    if pPr is not None:
        numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        return numPr is not None
    return False


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------
def extract_xlsx(filepath, sheet_name=None):
    """Extract table data from an XLSX file.

    Reads the specified sheet (or the active sheet) and returns it as a
    single table element with the first row as headers.
    """
    from openpyxl import load_workbook

    wb = load_workbook(filepath, read_only=True, data_only=True)
    if sheet_name:
        if sheet_name not in wb.sheetnames:
            print(f"  Sheet '{sheet_name}' not found. Available: {wb.sheetnames}")
            wb.close()
            return []
        ws = wb[sheet_name]
    else:
        ws = wb.active

    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append([str(c) if c is not None else "" for c in row])
    wb.close()

    if not rows:
        return []

    # Skip fully empty rows
    rows = [r for r in rows if any(cell.strip() for cell in r)]
    if len(rows) < 1:
        return []

    return [{
        "type": "table",
        "headers": rows[0],
        "rows": rows[1:],
    }]


# ---------------------------------------------------------------------------
# Post-processing
# ---------------------------------------------------------------------------
def _merge_paragraphs(elements):
    """Merge consecutive short paragraphs that likely belong together."""
    merged = []
    for el in elements:
        if (el["type"] == "paragraph"
                and merged
                and merged[-1]["type"] == "paragraph"
                and len(merged[-1]["text"]) < 80
                and not merged[-1]["text"].endswith(".")
                and not merged[-1]["text"].endswith(":")
                and len(el["text"]) < 200):
            merged[-1]["text"] += " " + el["text"]
        else:
            merged.append(el)
    return merged


# ---------------------------------------------------------------------------
# EditorJS conversion
# ---------------------------------------------------------------------------
def _make_id(prefix="blk"):
    """Generate a short unique block ID."""
    return f"{prefix}_{uuid.uuid4().hex[:8]}"


def to_editorjs(elements):
    """Convert normalized content elements to EditorJS blocks.

    Returns a list of block dicts ready for the Rattle API content endpoint.
    """
    blocks = []

    for el in elements:
        if el["type"] == "heading":
            blocks.append({
                "id": _make_id("h"),
                "type": "header",
                "data": {"text": el["text"], "level": el.get("level", 3)},
            })

        elif el["type"] == "paragraph":
            blocks.append({
                "id": _make_id("p"),
                "type": "paragraph",
                "data": {"text": el["text"]},
            })

        elif el["type"] == "table":
            content = [el["headers"]] + el.get("rows", [])
            blocks.append({
                "id": _make_id("tbl"),
                "type": "table",
                "data": {
                    "withHeadings": True,
                    "stretched": True,
                    "content": content,
                },
            })

        elif el["type"] == "list":
            items = [
                {"content": item, "items": [], "meta": {}}
                for item in el["items"]
            ]
            blocks.append({
                "id": _make_id("lst"),
                "type": "list",
                "data": {
                    "style": el.get("style", "unordered"),
                    "meta": {},
                    "items": items,
                },
            })

    return blocks


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Extract structured content from PDF, DOCX, or XLSX files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python tools/extract_content.py datasheet.pdf\n"
            "  python tools/extract_content.py datasheet.docx --format editorjs\n"
            "  python tools/extract_content.py specs.xlsx --sheet 'Sheet1'\n"
            "  python tools/extract_content.py datasheet.pdf -o content.json\n"
        ),
    )
    parser.add_argument("input_file", help="Path to the input file (.pdf, .docx, .xlsx)")
    parser.add_argument(
        "--format", choices=["json", "editorjs"], default="json",
        help="Output format: 'json' for normalized elements, 'editorjs' for API-ready blocks (default: json)",
    )
    parser.add_argument(
        "--sheet", default=None,
        help="Sheet name for XLSX files (default: active sheet)",
    )
    parser.add_argument(
        "-o", "--output", default=None,
        help="Output file path (default: print to stdout)",
    )

    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print(f"Error: File not found: {args.input_file}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(args.input_file)[1].lower()

    # Extract based on file type
    if ext == ".pdf":
        elements = extract_pdf(args.input_file)
    elif ext == ".docx":
        elements = extract_docx(args.input_file)
    elif ext in (".xlsx", ".xls"):
        elements = extract_xlsx(args.input_file, sheet_name=args.sheet)
    else:
        print(f"Error: Unsupported file type '{ext}'. Use .pdf, .docx, or .xlsx", file=sys.stderr)
        sys.exit(1)

    # Convert to EditorJS if requested
    if args.format == "editorjs":
        output = to_editorjs(elements)
    else:
        output = elements

    # Output
    result = json.dumps(output, indent=2, ensure_ascii=False)

    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(result)
        print(f"Wrote {len(output)} {'blocks' if args.format == 'editorjs' else 'elements'} to {args.output}")
    else:
        print(result)

    # Summary to stderr so it doesn't pollute piped output
    counts = {}
    for el in (output if args.format == "json" else elements):
        t = el.get("type", "unknown")
        counts[t] = counts.get(t, 0) + 1
    summary = ", ".join(f"{v} {k}{'s' if v > 1 else ''}" for k, v in sorted(counts.items()))
    print(f"\nExtracted: {summary}", file=sys.stderr)


if __name__ == "__main__":
    main()
